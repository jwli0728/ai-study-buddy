import os
import uuid as uuid_module
from pathlib import Path
from uuid import UUID

from sqlalchemy import select, func
from sqlalchemy.ext.asyncio import AsyncSession
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from pypdf import PdfReader
from docx import Document as DocxDocument

from app.db.models import Document, DocumentChunk
from app.schemas import DocumentResponse, DocumentListResponse, DocumentStatusResponse
from app.config import get_settings

settings = get_settings()


class DocumentService:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""],
        )
        self._embeddings = None

    @property
    def embeddings(self):
        if self._embeddings is None:
            self._embeddings = GoogleGenerativeAIEmbeddings(
                model=settings.EMBEDDING_MODEL,
                google_api_key=settings.GOOGLE_API_KEY,
            )
        return self._embeddings

    async def save_uploaded_file(
        self,
        db: AsyncSession,
        session_id: UUID,
        user_id: UUID,
        file_content: bytes,
        original_filename: str,
        mime_type: str,
    ) -> Document:
        """Save an uploaded file and create a document record."""
        # Generate unique filename
        file_ext = Path(original_filename).suffix
        unique_filename = f"{uuid_module.uuid4()}{file_ext}"

        # Ensure upload directory exists
        upload_dir = Path(settings.UPLOAD_DIR)
        upload_dir.mkdir(parents=True, exist_ok=True)

        file_path = upload_dir / unique_filename

        # Save file
        with open(file_path, "wb") as f:
            f.write(file_content)

        # Create document record
        document = Document(
            session_id=session_id,
            user_id=user_id,
            filename=unique_filename,
            original_filename=original_filename,
            file_path=str(file_path),
            file_size=len(file_content),
            mime_type=mime_type,
            processing_status="pending",
        )
        db.add(document)
        await db.commit()
        await db.refresh(document)
        return document

    async def process_document(self, db: AsyncSession, document_id: UUID) -> None:
        """Process a document: extract text, chunk, and create embeddings."""
        # Get document
        result = await db.execute(select(Document).where(Document.id == document_id))
        document = result.scalar_one_or_none()
        if not document:
            return

        try:
            # Update status to processing
            document.processing_status = "processing"
            await db.commit()

            # Extract text based on file type
            text = self._extract_text(document.file_path, document.mime_type)

            # Split into chunks
            chunks = self.text_splitter.split_text(text)

            if not chunks:
                document.processing_status = "completed"
                document.chunk_count = 0
                await db.commit()
                return

            # Generate embeddings
            embeddings = await self.embeddings.aembed_documents(chunks)

            # Store chunks with embeddings
            for i, (chunk_text, embedding) in enumerate(zip(chunks, embeddings)):
                chunk = DocumentChunk(
                    document_id=document.id,
                    session_id=document.session_id,
                    chunk_index=i,
                    content=chunk_text,
                    embedding=embedding,
                    metadata={"source": document.original_filename, "chunk_index": i},
                )
                db.add(chunk)

            # Update document status
            document.processing_status = "completed"
            document.chunk_count = len(chunks)
            await db.commit()

        except Exception as e:
            document.processing_status = "failed"
            document.processing_error = str(e)
            await db.commit()
            raise

    def _extract_text(self, file_path: str, mime_type: str) -> str:
        """Extract text from a file based on its type."""
        if mime_type == "application/pdf":
            return self._extract_pdf_text(file_path)
        elif mime_type in ["text/plain", "text/markdown"]:
            return self._extract_plain_text(file_path)
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return self._extract_docx_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {mime_type}")

    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from a PDF file."""
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return "\n\n".join(text_parts)

    def _extract_plain_text(self, file_path: str) -> str:
        """Extract text from a plain text file."""
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from a DOCX file."""
        doc = DocxDocument(file_path)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text_parts.append(paragraph.text)
        return "\n\n".join(text_parts)

    @staticmethod
    async def get_document(db: AsyncSession, document_id: UUID) -> Document | None:
        """Get a document by ID."""
        result = await db.execute(select(Document).where(Document.id == document_id))
        return result.scalar_one_or_none()

    @staticmethod
    async def get_document_for_user(
        db: AsyncSession, document_id: UUID, user_id: UUID
    ) -> Document | None:
        """Get a document by ID, verifying ownership."""
        result = await db.execute(
            select(Document).where(
                Document.id == document_id,
                Document.user_id == user_id,
            )
        )
        return result.scalar_one_or_none()

    @staticmethod
    async def list_documents(
        db: AsyncSession, session_id: UUID, skip: int = 0, limit: int = 50
    ) -> DocumentListResponse:
        """List all documents for a session."""
        query = select(Document).where(Document.session_id == session_id)

        # Get total count
        count_query = select(func.count()).select_from(query.subquery())
        total_result = await db.execute(count_query)
        total = total_result.scalar() or 0

        # Get paginated results
        query = query.order_by(Document.created_at.desc()).offset(skip).limit(limit)
        result = await db.execute(query)
        documents = result.scalars().all()

        return DocumentListResponse(
            documents=[DocumentResponse.model_validate(d) for d in documents],
            total=total,
        )

    @staticmethod
    async def delete_document(db: AsyncSession, document: Document) -> None:
        """Delete a document and its file."""
        # Delete file
        if os.path.exists(document.file_path):
            os.remove(document.file_path)

        # Delete record (chunks will cascade)
        await db.delete(document)
        await db.commit()

    @staticmethod
    async def get_document_status(
        db: AsyncSession, document_id: UUID
    ) -> DocumentStatusResponse | None:
        """Get the processing status of a document."""
        result = await db.execute(select(Document).where(Document.id == document_id))
        document = result.scalar_one_or_none()
        if not document:
            return None
        return DocumentStatusResponse(
            id=document.id,
            processing_status=document.processing_status,
            processing_error=document.processing_error,
            chunk_count=document.chunk_count,
        )
